from tkinter import Tk
from tkinter import filedialog
from docx import Document
from docx.shared import Pt, Length, RGBColor
import textwrap
import requests
from transformers import pipeline
import string
from sbert_punc_case_ru import SbertPuncCase
from pyaspeller import YandexSpeller
import urllib
import json
from urllib.request import urlopen

# для загрузки sbert_punc_case_ru
# pip install git+https://huggingface.co/kontur-ai/sbert_punc_case_ru

# Текст: Заголовок Стиль: Заголовок без нумерации
# Текст: Подзаголовок Стиль: Heading 2
# Текст: Пункт Стиль: Heading 3
# Текст: Подпункт Стиль: Heading 4
# Текст: Обычный текст Стиль: Normal

# список союзов и предлогов
prepos = [
    "то", "это", "и", "да", "ни", "как", "так", "нетолько", "но",
    "также", "тоже", "или", "либо", "толи", "не", "а", "но", "да",
    "за", "однако", "что", "бы",  "будто", "словно", "точно", "чем", "тем", "когда", "едва", "только",
    "пока", "прежде", "чем", "в", "ли", "после", "того", "чтобы", "для","того", "чтобы", "ради", "дабы",
    "с", "тем", "потому", "оттого что", "ибо", "если",
    "коли", "ежели", "хотя", "хоть", "несмотря", "на", "пусть", "пускай",
]

def remove_repeated_words(text:str)->str:
    words = text.split()
    new_words = []
    prev_word = None
    for word in words:
        if word != prev_word:
            new_words.append(word)
        prev_word = word
    return ' '.join(new_words)

def correct(text: str, style: str)->str:
    speller = YandexSpeller(lang="ru",
                            ignore_digits=True,
                            ignore_urls=True,
                            find_repeat_words=True
                            )

    fixed = []
    text = remove_repeated_words(text)
    # удаляем все знаки препинания и приводим к нижнему регистру текст параграфа только если это обычный текст
    if style == 'Normal':
        text = text.translate(str.maketrans('', '', string.punctuation)).lower().split()
    else:
        text = text.split()

    for error in text:
        fixed.append(speller.spelled(error))
    return ' '.join(fixed)


def dockLoad():
    Tk().withdraw()
    filepath = filedialog.askopenfilename()
    # filepath = 'C:/Users/Администратор/Desktop/queta.docx'
    if filepath:
        doc = Document(filepath)

        # проходимся по все параграфам
        for para in doc.paragraphs:
            # получаем прогоны параграфа, в них хранятся все свойства текста
            run = para.add_run()

            # Настройки параграфа
            para.line_spacing_rule = 1
            para.paragraph_format.space_after = 0
            para.paragraph_format.space_before = 0

            # исправляем орфографические ошибки
            para.text = remove_repeated_words(correct(para.text, para.style.name))

            # Если обычный текст, то кроме исправления стиля так же расставляем знаки пунктуации
            if para.style.name == 'Normal':
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11)
                max_width = 0

                # импортируем модель и расставляем знаки препинания с помощью модели
                model = SbertPuncCase()
                para.text = model.punctuate(para.text)

            # Всему, что не обычный текст, просто меняем стиль
            if para.style.name == 'Заголовок без нумерации':
                run.font.name = 'Segoe UI'
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255,78,19)


            if para.style.name == 'Heading 2':
                run.font.name = 'Segoe UI'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            if para.style.name == 'Heading 3':
                run.font.name = 'Segoe UI'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            if para.style.name == 'Heading 4':
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        filename = filedialog.asksaveasfilename(defaultextension=".docx")
        if filename:
            doc.save(filename)

